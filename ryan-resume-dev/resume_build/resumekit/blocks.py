"""One renderer per content block type.

The YAML content files are a list of sections, each with a `type` that names a
function here. Adding a section type means adding a function and registering it
in RENDERERS — nothing else changes.
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from . import brand
from .brand import METRICS as M
from .docx_helpers import (
    keep_lines_together,
    keep_with_next,
    set_paragraph_border,
    set_table_cell_margins,
    strip_table_borders,
)


# --------------------------------------------------------------------------
# Shared pieces
# --------------------------------------------------------------------------

def section_heading(doc, text: str, note: str | None = None):
    """Heading + coral underline. This is coral slot 3 of 3 (brand.CORAL_RATION)."""
    p = doc.add_paragraph(style="WP Section")
    p.add_run(text)
    if note:
        # Parenthetical qualifier, e.g. the year on "Active Development Areas".
        run = p.add_run(f"  {note}")
        run.bold = False
        run.font.size = Pt(M.meta_pt)
        run.font.color.rgb = RGBColor.from_string(brand.INK_45)
    set_paragraph_border(
        p, "bottom", brand.CORAL, M.section_rule_eighths, M.section_rule_gap_pt
    )
    keep_with_next(p)
    return p


def _bullet(doc, text: str):
    p = doc.add_paragraph(style="WP Bullet")
    glyph = p.add_run(f"{M.bullet_glyph}\t")
    glyph.font.color.rgb = RGBColor.from_string(brand.INK_45)
    p.add_run(text)
    keep_lines_together(p)
    return p


def _body(doc, text: str):
    p = doc.add_paragraph(style="WP Body")
    p.add_run(text)
    return p


# --------------------------------------------------------------------------
# Block types
# --------------------------------------------------------------------------

def render_prose(doc, sec: dict) -> None:
    """type: prose — a heading and one or more paragraphs."""
    section_heading(doc, sec["title"], sec.get("note"))
    for para in sec["paragraphs"]:
        _body(doc, para)


def render_grid(doc, sec: dict) -> None:
    """type: grid — the Core Expertise table. Borderless two-column layout."""
    section_heading(doc, sec["title"], sec.get("note"))

    items = sec["items"]
    cols = int(sec.get("columns", 2))
    rows = (len(items) + cols - 1) // cols

    table = doc.add_table(rows=rows, cols=cols)
    strip_table_borders(table)
    set_table_cell_margins(table, top=0.0, start=0.0, bottom=0.04, end=0.12)
    table.autofit = False

    usable = 7.5 - 0.02          # a hair under the text column so Word can't wrap
    for col in table.columns:
        col.width = Inches(usable / cols)

    for i, item in enumerate(items):
        cell = table.cell(i // cols, i % cols)
        # Every cell starts with one empty paragraph; reuse it for the label.
        label = cell.paragraphs[0]
        label.style = doc.styles["WP Skill Label"]
        label.add_run(item["label"])

        listing = cell.add_paragraph(style="WP Skill List")
        listing.add_run(M.skill_sep.join(item["skills"]))

    # Word's default table cell width can override the column width; force it.
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(usable / cols)


def render_experience(doc, sec: dict) -> None:
    """type: experience — role title, mono date line, bullets.

    Optional `closing:` adds one body paragraph after the last role. Added v2.0
    so the merged music section can end on a single credentials line instead of
    needing a second section heading for three facts.
    """
    section_heading(doc, sec["title"], sec.get("note"))
    for role in sec["roles"]:
        title = doc.add_paragraph(style="WP Role")
        headline = role["title"]
        if role.get("org"):
            headline = f"{headline} — {role['org']}"
        title.add_run(headline)
        keep_with_next(title)

        if role.get("dates"):
            meta = doc.add_paragraph(style="WP Role Meta")
            meta.add_run(role["dates"])
            keep_with_next(meta)

        for text in role.get("bullets", []):
            _bullet(doc, text)

    if sec.get("closing"):
        _body(doc, sec["closing"])


def render_bullets(doc, sec: dict) -> None:
    """type: bullets — a heading and a flat bullet list."""
    section_heading(doc, sec["title"], sec.get("note"))
    for text in sec["bullets"]:
        _bullet(doc, text)


def render_lines(doc, sec: dict) -> None:
    """type: lines — a heading and one or more single-line facts (Education)."""
    section_heading(doc, sec["title"], sec.get("note"))
    for text in sec["lines"]:
        _body(doc, text)


def render_projects(doc, sec: dict) -> None:
    """type: projects — named project, then a descriptive paragraph."""
    section_heading(doc, sec["title"], sec.get("note"))
    for item in sec["items"]:
        name = doc.add_paragraph(style="WP Project")
        name.add_run(item["name"])
        keep_with_next(name)
        _body(doc, item["body"])


RENDERERS = {
    "prose": render_prose,
    "grid": render_grid,
    "experience": render_experience,
    "bullets": render_bullets,
    "lines": render_lines,
    "projects": render_projects,
}
