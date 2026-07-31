"""Content dict -> .docx.

build_resume() is the whole public surface. Everything it does is either a
brand token (brand.py), a style (styles.py), or a block renderer (blocks.py).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, RGBColor

from . import brand
from .blocks import RENDERERS
from .brand import METRICS as M
from .docx_helpers import (
    add_field,
    set_char_spacing,
    set_paragraph_border,
    set_tab_stops,
    style_field_runs,
)
from .styles import build_styles


# --------------------------------------------------------------------------
# Page setup
# --------------------------------------------------------------------------

def _setup_page(doc) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(M.margin_in))
    sec.footer_distance = Inches(M.footer_distance_in)
    # The design is drawn on a 7.5in text column — the core-expertise table and
    # the footer's right tab stop are both measured off it. Spec §8 step 1.
    assert sec.left_margin == sec.right_margin == Inches(M.margin_in)


# --------------------------------------------------------------------------
# Header — plain text, always (spec §9)
#
# v2.2 removed the image banner. It was the default through v2.1, and a text
# fallback existed behind --header text for online applications. Extracting the
# text from a v2.1 build showed why that split was the wrong way round: the
# document's first parseable string was "Professional Summary" — the name,
# email, LinkedIn, GitHub, and location existed only as pixels. A parser built
# a candidate record with no name and no way to make contact. A résumé that is
# only sometimes parseable is one wrong export away from being unparseable, so
# there is now one header and it is text.
# --------------------------------------------------------------------------

def _add_text_header(doc, variant: str) -> None:
    """Real, selectable, parseable text carrying the banner's design cues —
    same navy, same coral rule, same mono contact line."""
    c = brand.CONTACT

    name = doc.add_paragraph(style="WP ATS Name")
    name.add_run(c["name"])

    role = doc.add_paragraph(style="WP ATS Role")
    role.add_run(brand.ROLE_LINES[variant])

    contact = doc.add_paragraph(style="WP ATS Contact")
    contact.add_run(
        f"{c['email']} · {c['linkedin']} · {c['github']} · {c['location']}"
    )
    # Coral slot 1 of 3 (brand.CORAL_RATION).
    set_paragraph_border(
        contact, "bottom", brand.CORAL, M.ats_rule_eighths, M.section_rule_gap_pt
    )


# --------------------------------------------------------------------------
# Footer — native text so the page number is a live field (spec §7)
# --------------------------------------------------------------------------

def _add_footer(doc, fonts: dict, org: str) -> None:
    c = brand.CONTACT
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0]
    p.style = doc.styles["WP Footer"]
    set_tab_stops(p, [(M.footer_tab_in, "right")])
    # Coral slot 2 of 3.
    set_paragraph_border(
        p, "top", brand.CORAL, M.footer_rule_eighths, M.footer_rule_gap_pt
    )

    p.add_run(f"{c['name']} · {c['email']} · {org}\t")
    p.add_run("Page ")

    mono = fonts["mono"]
    for instruction in ("PAGE", "NUMPAGES"):
        if instruction == "NUMPAGES":
            p.add_run(" of ")
        runs = add_field(p, instruction, placeholder="1")
        style_field_runs(
            runs, mono, M.footer_pt, brand.INK_45,
            tracking_pt=M.footer_tracking_pt, caps=True,
        )

    # The static half of the line is INK_70; the page counter is INK_45 so the
    # eye lands on the identity, not the pagination.
    for run in p.runs[:2]:
        run.font.color.rgb = RGBColor.from_string(brand.INK_70)


# --------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------

def build_resume(content: dict, out_path: Path, *,
                 fonts: dict | None = None,
                 density: str | None = None,
                 version: str = "") -> Path:
    """content: a parsed YAML content file. Keyword args override its `meta`."""
    meta = content["meta"]
    variant = meta["variant"]
    fonts = fonts or brand.FONT_SETS["brand"]

    # Must happen before build_styles(), which reads METRICS.
    brand.apply_density(density or meta.get("density", "default"))

    if variant not in brand.ROLE_LINES:
        raise ValueError(f"unknown variant {variant!r}")

    doc = Document()
    _setup_page(doc)
    build_styles(doc, fonts)
    _add_text_header(doc, variant)

    for sec in content["sections"]:
        kind = sec["type"]
        if kind not in RENDERERS:
            raise ValueError(
                f"unknown section type {kind!r}; known: {sorted(RENDERERS)}"
            )
        RENDERERS[kind](doc, sec)

    _add_footer(doc, fonts, meta.get("footer_org", brand.CONTACT["org"]))

    props = doc.core_properties
    props.author = brand.CONTACT["name"]
    props.title = meta.get("title", f"{brand.CONTACT['name']} — Résumé")
    props.subject = meta.get("subject", "")
    # The version lives in the document properties as well as the filename, so a
    # --release copy (which drops it from the name) can still be traced back.
    if version:
        props.version = version
    props.comments = (
        f"Generated by ryan-resume-dev/resume_build"
        + (f" v{version}" if version else "")
        + f" from {meta.get('source', 'content')} — do not hand-edit; edit the YAML."
    )

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
